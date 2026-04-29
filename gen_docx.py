import sys
sys.path.insert(0, r'D:\Fish\桌面\Note\my-obsidian')

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_LINE_SPACING
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'python-docx', '-q'])
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_LINE_SPACING

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(10.5)  # 五号 = 10.5pt
style.element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '宋体')

# Set line spacing to 18pt
pf = style.paragraph_format
pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
pf.line_spacing = Pt(18)
pf.space_before = Pt(0)
pf.space_after = Pt(0)

# Set margins
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

title = "《感性工学》课程感想"

paragraphs_text = [
    "坦白说，在上这门课之前，我对《感性工学》的想象仅限于"教我们怎么把东西做得更好看、更让人喜欢"的层面。作为一个工业设计专业的学生，我习惯了画草图、建模型，觉得自己离"科学计算情感"这件事非常遥远。然而，在为期数周的课程以及最后一次的总结回顾后，我发现自己之前的认知太浅薄了。感性工学并非关于"好不好看"的玄学，而是一场试图在主观心理与客观物理之间搭建桥梁的勇敢尝试。",
    "一、 课程核心认识与引发的迷思",
    "课程最核心的逻辑，就是将模糊的感性需求定量化。通过语义差分法、因子分析等手段，把"高级感"、"温暖"这类虚无缥缈的形容词，翻译成具体的物理设计参数。这让我第一次意识到，原来我们做设计时的那些直觉，是可以被拆解、被追溯的。在参与"数量化一类"的造型实验和"温觉实验"时，我亲身经历了从筛选词汇，到处理SPSS数据，最后建立线性回归模型的全过程。看着数学公式推导出了与主观感受高度相关的结论，那一瞬间确实有种"数据掌握真理"的震撼感。",
    "然而，随着课程的深入，尤其是最后几节课提出的那些核心反思，我开始感到困惑和不安。老师PPT上那句"模型输出的是噪点，而不是灵魂"像一根针一样扎进了我的脑海。当模型因为追求统计学的"安全"，而陷入"平均值陷阱"时，设计师的创意是否也被绑架了？当算法试图用冰冷的公式去定义"爱"时，那个在京都长大的设计师与在台中长大的设计师内心深处唤醒的不同意象，又该何去何从？我认识到，感性工学的强大之处在于它能提供理性的脚手架，但它的边界也很清晰：如果我们迷信数据，设计师就可能沦为"数据解释者"，而丢失了那种在某一瞬间突然通晓"这个形态应该像水滴一样"的、无法言说的感性跃迁能力。",
    "这引出了一个我至今仍在思考的问题：当感性工学的数据结论与我的审美直觉发生剧烈冲突时，我该怎么办？ 是相信那套严谨的、由样本推导出的数学模型，还是坚持那个只有我自己能感受到的、带着个人所有生命体验的直觉判断？这门课没有给我标准答案，但它教会了我如何清醒地看待这种冲突。",
    "二、 对未来应用与发展的思考",
    "关于感性工学的未来，我既兴奋又抱有警惕。教学大纲里提到的人工智能与眼动追踪、大数据分析的结合，无疑是未来的趋势。正如课堂上探讨的"人工智能如何介入感性工学"，我觉得AI可以成为感性工学进化的催化剂。它有能力处理更海量的、动态的数据，可以在一定程度上缓解老师提到过的"局限性二"——数据驱动的时效性滞后问题。比如，利用深度学习去实时抓取社交媒体上的审美风向，预测"审美疲劳"的临界点，这会让感性工学不再只是一张滞后的"静态快照"。",
    "但是，AI的介入也加深了另一层忧虑。当机器能无限逼近甚至预测人类的平均情感反应时，设计师的"灵魂"是否会被进一步挤压？授课老师在总结中给出的那个定位我非常认同：感性工学应该是"辅助手段而非终极目的"。不论技术如何发展，设计师的情感决策核心地位不应被撼动。未来的感性工学，或许不会是通过计算来完全复刻"爱"，而是利用多维整合——比如用眼动数据验证视觉焦点，用生理信号印证情感波动，用大数据描绘文化背景——来为设计师打开一扇更清晰地看见用户的窗户，但窗外的那片风景，终究需要设计师用自己带有文化体温的目光去解读。",
    "三、 课程兴趣点与改进建议",
    "在课程各部分内容中，我最感兴趣的是关于"材料温觉"的实验部分。以往我们选材更多地依赖经验式的"手感"，而温觉实验让我看到，人对材料"冷暖"的感知，背后有着严谨的物理学和生理学基础。将手触碰不同材质那一瞬间的温差变化，转化为可量化的图表，这种将极度私人的体感公开化、逻辑化的过程，充满了理性的美感。",
    "对于课程，我有以下两点小小的改进建议：",
    "第一，关于局限性的探讨可以前置。我是到课程末尾听到那页总结时才豁然开朗的，如果能在课程的早中期，当我们刚开始学习方法时，就同步引入"感性工学的局限性"讨论，或许大家在实验过程中会带着更批判性的眼光看待数据，而不是盲目追求模型的好看。",
    "第二，希望能增加一些利用人工智能工具辅助分析的实践环节。课程大纲里提到了智谱清言等工具的应用，但在实际教学中，我们很多同学对如何使用DeepSeek等工具高效检索文献、辅助解读SPSS数据结果还不太熟练。如果能有一些更具体的操作引导，把这一块真正落实下来，或许能让我们的数据处理能力和效率都得到进一步的提升。",
    "总而言之，这门课给予我的远不止两种实验方法。它让我学会尊重数据，更让我学会警惕数据。设计不仅是数据的映射，更是文化与感性的共鸣。在走向高度自动化的数据时代，守护好内心深处那份脆弱的、非理性的、但足够珍贵的感性直觉，或许就是这门课留给我最深的课题。"
]

# Title - centered, bold
p_title = doc.add_paragraph()
p_title.alignment = 1  # CENTER
p_title.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
p_title.paragraph_format.line_spacing = Pt(18)
run = p_title.add_run(title)
run.font.name = '宋体'
run.font.size = Pt(10.5)
run.font.bold = True
run.element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '宋体')

# Body paragraphs
for text in paragraphs_text:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.first_line_indent = Pt(10.5 * 2)  # 2 char indent
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(10.5)
    run.element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '宋体')

# Remove first line indent for section headings
for para in doc.paragraphs:
    text = para.text.strip()
    if text.startswith("一、") or text.startswith("二、") or text.startswith("三、"):
        para.paragraph_format.first_line_indent = Pt(0)

output_path = r'D:\Fish\桌面\Note\my-obsidian\231301-3123002261-余卓成.docx'
doc.save(output_path)
print(f"Saved to: {output_path}")
